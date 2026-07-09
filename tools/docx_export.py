import io
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "Calibri"
NAME_SIZE = Pt(22)
CONTACT_SIZE = Pt(10)
SECTION_SIZE = Pt(11)
BODY_SIZE = Pt(10.5)
BULLET_SIZE = Pt(10.5)


# ---------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------

def _set_font(run, size, bold=False, italic=False, color=None):
    run.font.name = FONT_NAME
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def _add_bottom_border(paragraph):
    """Adds a thin bottom border line under a paragraph (section headers)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_paragraph_spacing(paragraph, before=0, after=2):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)


def _add_runs_with_inline_bold(paragraph, text, base_size=BODY_SIZE, base_bold=False, base_italic=False):
    """
    Parses **bold** and *italic* markers in text and adds Word runs
    with correct formatting. Handles nested or sequential markers.
    """
    # Split on **...** first, then on *...*
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            _set_font(run, base_size, bold=True, italic=base_italic)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            _set_font(run, base_size, bold=base_bold, italic=True)
        elif part:
            run = paragraph.add_run(part)
            _set_font(run, base_size, bold=base_bold, italic=base_italic)


# ---------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------

def _add_name_header(doc, name, contact):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p, before=0, after=2)
    run = p.add_run(name.strip())
    _set_font(run, NAME_SIZE, bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p2, before=0, after=4)
    run2 = p2.add_run(contact.strip())
    _set_font(run2, CONTACT_SIZE)
    _add_bottom_border(p2)


def _add_section_header(doc, title):
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=6, after=2)
    run = p.add_run(title.upper())
    _set_font(run, SECTION_SIZE, bold=True)
    _add_bottom_border(p)
    return p


def _add_job_title(doc, text):
    """Bold job title line: Company | Role | Location | Dates"""
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=4, after=1)
    # Strip leading markdown bold markers if present
    clean = re.sub(r"^\*+|\*+$", "", text).strip()
    _add_runs_with_inline_bold(p, clean, base_size=BODY_SIZE, base_bold=True)


def _add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    _set_paragraph_spacing(p, before=1, after=1)
    fmt = p.paragraph_format
    fmt.left_indent = Inches(0.2)
    fmt.first_line_indent = Inches(-0.2)
    # Replace markdown list markers
    clean = re.sub(r"^[-•*]\s*", "", text).strip()
    _add_runs_with_inline_bold(p, clean, base_size=BULLET_SIZE)


def _add_body_line(doc, text, italic=False):
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=1, after=1)
    _add_runs_with_inline_bold(p, text.strip(), base_size=BODY_SIZE, base_italic=italic)


def _add_skills_line(doc, text):
    """Renders 'Label: skill · skill · skill' with the label bolded."""
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=1, after=1)
    # Bold the label (text before first colon)
    if ":" in text:
        label, rest = text.split(":", 1)
        run = p.add_run(label.strip() + ": ")
        _set_font(run, BODY_SIZE, bold=True)
        run2 = p.add_run(rest.strip())
        _set_font(run2, BODY_SIZE)
    else:
        _add_runs_with_inline_bold(p, text.strip(), base_size=BODY_SIZE)


# ---------------------------------------------------------------
# Main parser
# ---------------------------------------------------------------

def _is_section_header(line):
    """Detect section header patterns from the AI output."""
    clean = re.sub(r"[#\-*`]", "", line).strip()
    return (
        clean.isupper()
        and len(clean) > 2
        and len(clean) < 60
        and not clean.startswith("•")
    )


def _is_job_title(line):
    """Lines that look like 'Company — Role | Location | Dates' or similar."""
    stripped = line.strip()
    # Bold markdown job title or line with | pipe separators and a date pattern
    has_pipes = stripped.count("|") >= 2
    has_date = bool(re.search(r"\b(20\d\d|Present)\b", stripped))
    is_bold_line = stripped.startswith("**") and stripped.endswith("**")
    return (has_pipes and has_date) or is_bold_line


def _is_bullet(line):
    return bool(re.match(r"^\s*[-•*]\s+", line))


def _is_divider(line):
    return bool(re.match(r"^[-_*]{3,}\s*$", line.strip()))


def _is_italic_line(line):
    stripped = line.strip()
    return stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**")


KNOWN_SECTION_WORDS = {
    "EXPERIENCE", "WORK EXPERIENCE", "SKILLS", "SUMMARY", "EDUCATION",
    "AWARDS", "CERTIFICATIONS", "PROJECTS"
}


def _resolve_contact_placeholders(text: str) -> str:
    """Replace [PHONE] and [EMAIL] with runtime values from secrets/env."""
    try:
        import streamlit as st
        phone = st.secrets.get("CONTACT_PHONE", "") or os.getenv("CONTACT_PHONE", "[PHONE]")
        email = st.secrets.get("CONTACT_EMAIL", "") or os.getenv("CONTACT_EMAIL", "[EMAIL]")
    except Exception:
        phone = os.getenv("CONTACT_PHONE", "[PHONE]")
        email = os.getenv("CONTACT_EMAIL", "[EMAIL]")
    return text.replace("[PHONE]", phone).replace("[EMAIL]", email)


def resume_to_docx(resume_text: str, name: str = "SHRUTHI TADAMIRI",
                   contact: str = "[PHONE] | [EMAIL] | linkedin.com/in/shruthi-tadamiri | Boston, MA") -> bytes:
    """
    Converts the tailored resume text into a formatted .docx matching the
    reference resume style: centered name header, section headers with bottom
    borders, bold job titles, indented bullets with inline bold, skills lines
    with bold labels.
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Remove default paragraph spacing from Normal style
    doc.styles["Normal"].paragraph_format.space_after = Pt(0)
    doc.styles["Normal"].font.name = FONT_NAME

    _add_name_header(doc, name, _resolve_contact_placeholders(contact))

    lines = resume_text.split("\n")

    # Strip the leading name/contact block from the LLM output — it's already
    # rendered by _add_name_header above. Skip until we hit the first real
    # section header (SUMMARY, SKILLS, EXPERIENCE, etc.).
    first_section_idx = 0
    for i, line in enumerate(lines):
        clean_test = re.sub(r"[#\-*`]", "", line).strip()
        label_test = re.sub(r"[*_`]", "", clean_test).strip()
        if label_test.upper() in KNOWN_SECTION_WORDS or (
            clean_test.isupper() and 2 < len(clean_test) < 60 and not clean_test.startswith("•")
        ):
            first_section_idx = i
            break
    lines = lines[first_section_idx:]

    for line in lines:
        stripped = line.strip()

        if not stripped:
            continue

        if _is_divider(stripped):
            continue

        # Strip leading # from markdown headers
        clean = re.sub(r"^#+\s*", "", stripped)

        # Detect section headers
        label = re.sub(r"[*_`]", "", clean).strip()
        if label.upper() in KNOWN_SECTION_WORDS or _is_section_header(clean):
            _add_section_header(doc, label)
            continue

        # Bullets
        if _is_bullet(stripped):
            _add_bullet(doc, stripped)
            continue

        # Italic award/recognition lines
        if _is_italic_line(stripped):
            _add_body_line(doc, stripped.strip("*").strip(), italic=True)
            continue

        # Job title lines
        if _is_job_title(clean):
            _add_job_title(doc, clean)
            continue

        # Skills lines (contain ":" and "·" or are in a skills section)
        if ":" in clean and ("·" in clean or "," in clean):
            _add_skills_line(doc, clean)
            continue

        # Default: body line
        _add_body_line(doc, clean)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
